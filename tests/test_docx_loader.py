from docx import Document

from rag.loaders.docx_loader import DocxLoader


def test_docx_loader_preserves_heading_context_and_tables(tmp_path):
    file_path = tmp_path / "policy.docx"
    document = Document()
    document.add_heading("Cards", level=1)
    document.add_heading("Disputes", level=2)
    document.add_paragraph("Collect the transaction reference.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Issue"
    table.cell(0, 1).text = "SLA"
    table.cell(1, 0).text = "Unauthorized transaction"
    table.cell(1, 1).text = "3 working days"
    document.add_heading("Escalation", level=2)
    document.add_paragraph("Escalate unresolved cases.")
    document.save(file_path)

    chunks = DocxLoader().load(str(file_path))

    assert len(chunks) == 2
    assert "Cards\nDisputes" in chunks[0].text
    assert "Issue | SLA" in chunks[0].text
    assert "Unauthorized transaction | 3 working days" in chunks[0].text
    assert chunks[1].text.startswith("Cards\nEscalation")
